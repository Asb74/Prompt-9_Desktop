import logging
from pathlib import Path

from docx import Document
from pypdf import PdfReader

from src.services.table_loader import TableLoader


class DocumentLoader:
    CSV_MAX_ROWS = 200
    XLSX_MAX_ROWS_PER_SHEET = 300
    XLS_MAX_ROWS_PER_SHEET = 300
    INTERMEDIATE_CHAR_LIMIT = 120000

    def __init__(self) -> None:
        self.logger = logging.getLogger(__name__)
        self.table_loader = TableLoader()

    def extract_text(self, path: str) -> str:
        file_path = Path(path)
        extension = file_path.suffix.lower()

        if extension == ".txt":
            return self._extract_txt(file_path)
        if extension == ".csv":
            return self._extract_csv(file_path)
        if extension == ".pdf":
            return self._extract_pdf(file_path)
        if extension == ".docx":
            return self._extract_docx(file_path)
        if extension == ".doc":
            return self._extract_doc(file_path)
        if extension in {".xlsx", ".xls"}:
            return self._extract_spreadsheet(file_path)

        raise ValueError(f"Formato de archivo no soportado: {extension}")

    def _extract_txt(self, file_path: Path) -> str:
        for encoding, kwargs in (("utf-8", {}), ("latin-1", {}), ("utf-8", {"errors": "replace"})):
            try:
                with file_path.open("r", encoding=encoding, **kwargs) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        return ""

    def _extract_csv(self, file_path: Path) -> str:
        tables = self.table_loader.load_tables(str(file_path))
        if not tables:
            return ""

        preview = self.table_loader.table_to_preview_text(tables[0], max_rows=self.CSV_MAX_ROWS)
        return preview[: self.INTERMEDIATE_CHAR_LIMIT]

    def _extract_pdf(self, file_path: Path) -> str:
        reader = PdfReader(str(file_path))
        chunks: list[str] = []
        for idx, page in enumerate(reader.pages, start=1):
            try:
                page_text = page.extract_text() or ""
            except Exception:
                self.logger.warning("No se pudo extraer texto de página %s en PDF %s", idx, file_path.name)
                continue
            if page_text.strip():
                chunks.append(page_text)
            if sum(len(x) for x in chunks) >= self.INTERMEDIATE_CHAR_LIMIT:
                break
        return "\n\n".join(chunks)

    def _extract_docx(self, file_path: Path) -> str:
        doc = Document(str(file_path))
        parts: list[str] = []
        for paragraph in doc.paragraphs:
            text = (paragraph.text or "").strip()
            if text:
                parts.append(text)
        for table in doc.tables:
            for row in table.rows:
                cells = [((cell.text or "").strip()) for cell in row.cells]
                if any(cells):
                    parts.append("\t".join(cells))
                if sum(len(x) for x in parts) >= self.INTERMEDIATE_CHAR_LIMIT:
                    break
        return "\n".join(parts)


    def _extract_doc(self, file_path: Path) -> str:
        self.logger.info("Archivo .doc detectado: %s", file_path.name)
        converted_path = self.convert_doc_to_docx(file_path)
        if converted_path is not None and converted_path.exists():
            try:
                return self._extract_docx(converted_path)
            finally:
                if converted_path != file_path and converted_path.exists():
                    try:
                        converted_path.unlink()
                    except OSError:
                        self.logger.warning("No se pudo eliminar archivo temporal convertido: %s", converted_path)

        raise ValueError(
            "El formato .doc antiguo está permitido, pero requiere conversión previa a .docx para extraer texto correctamente."
        )

    def convert_doc_to_docx(self, file_path: Path) -> Path | None:
        """Punto de extensión opcional para convertir .doc a .docx en entornos con Word/LibreOffice."""
        self.logger.info("Conversión .doc -> .docx no configurada en este entorno: %s", file_path.name)
        return None

    def _extract_spreadsheet(self, file_path: Path) -> str:
        tables = self.table_loader.load_tables(str(file_path))
        parts: list[str] = []

        for table in tables:
            parts.append(self.table_loader.table_to_preview_text(table, max_rows=self.XLSX_MAX_ROWS_PER_SHEET))
            parts.append("")
            if sum(len(x) for x in parts) >= self.INTERMEDIATE_CHAR_LIMIT:
                break

        return "\n".join(parts)[: self.INTERMEDIATE_CHAR_LIMIT]
